import os
from docx import Document as DocxDocument
from app.models.document import Document
from datetime import datetime


class ExportServiceMock:

    def export(self, data: Document, output_dir: str):
        print("[ExportService] Exportando conteúdo...")

        os.makedirs(output_dir, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = os.path.join(output_dir, f"output_{timestamp}.txt")

        with open(output_file, "w", encoding="utf-8") as f:
            for paragraph in data.paragraphs:
                f.write(paragraph.text + "\n\n")

        print(f"[ExportService] Arquivo gerado em: {output_file}")

        # ✔ CORREÇÃO: retornar caminho do arquivo
        return output_file


# app/services/export_service.py
class ExportService:

    def export(self, data: Document, output_dir: str, format: str = "txt", segments=None):
        print("[ExportService] Exportando conteúdo...")

        os.makedirs(output_dir, exist_ok=True)

        # ✔ CORREÇÃO: retornar resultado das funções internas
        if format == "txt":
            return self._export_txt(data, output_dir)

        elif format == "docx":
            return self._export_docx(data, output_dir)

        elif format == "srt":
            return self._export_srt(output_dir, segments)

        else:
            raise ValueError(f"Formato não suportado: {format}")

    # =========================
    # TXT
    # =========================
    def _export_txt(self, data: Document, output_dir: str):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = os.path.join(output_dir, f"output_{timestamp}.txt")

        with open(output_file, "w", encoding="utf-8") as f:
            for paragraph in data.paragraphs:
                f.write(paragraph.text + "\n\n")

        print(f"[ExportService] TXT gerado em: {output_file}")

        # ✔ CORREÇÃO: retornar caminho do arquivo
        return output_file

    # =========================
    # DOCX
    # =========================
    def _export_docx(self, data: Document, output_dir: str):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = os.path.join(output_dir, f"transcricao_{timestamp}.docx")

        doc = DocxDocument()

        # Título (garantido)
        title = getattr(data, "title", "Transcrição")
        doc.add_heading(title, level=1)

        # Parágrafos
        for paragraph in data.paragraphs:
            p = doc.add_paragraph(paragraph.text)
            p.paragraph_format.space_after = 12

        doc.save(output_file)

        print(f"[ExportService] DOCX gerado em: {output_file}")

        # ✔ CORREÇÃO: retornar caminho do arquivo
        return output_file

    # =========================
    # SRT
    # =========================
    def _export_srt(self, output_dir: str, segments):
        if not segments:
            raise ValueError("Segments são necessários para exportação SRT")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = os.path.join(output_dir, f"transcricao_{timestamp}.srt")

        with open(output_file, "w", encoding="utf-8") as f:
            for i, seg in enumerate(segments, start=1):
                start = self._format_timestamp(seg.start)
                end = self._format_timestamp(seg.end)
                text = seg.text.strip()

                f.write(f"{i}\n")
                f.write(f"{start} --> {end}\n")
                f.write(f"{text}\n\n")

        print(f"[ExportService] SRT gerado em: {output_file}")

        return output_file

    def _format_timestamp(self, seconds: float) -> str:
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int(round((seconds % 1) * 1000))

        return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"